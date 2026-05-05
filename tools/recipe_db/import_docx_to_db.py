import argparse
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


UTC = timezone.utc


def now_utc_iso() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat()


def canonicalize_name(name: str) -> str:
    name = re.sub(r"\s+", " ", name.strip())
    name = re.sub(r"^\d+\.\s*", "", name)
    return name.lower()


def is_numbered_title_line(text: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+\S", text))


def strip_bullet_prefix(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^[•\-\u2022]\s*", "", text)
    return text.strip()


@dataclass(frozen=True)
class DishRow:
    display_name: str
    canonical_name: str
    category: str
    subcategory: str | None
    cuisine_group: str | None
    source_order: int


def extract_seafood(doc_path: Path) -> list[DishRow]:
    doc = Document(str(doc_path))
    cuisine_group: str | None = None
    out: list[DishRow] = []
    order = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = getattr(p.style, "name", "") or ""
        if style.startswith("Heading") and not is_numbered_title_line(text):
            cuisine_group = text
            continue
        if is_numbered_title_line(text):
            order += 1
            name = re.sub(r"^\s*\d+\.\s*", "", text).strip()
            out.append(
                DishRow(
                    display_name=name,
                    canonical_name=canonicalize_name(name),
                    category="seafood",
                    subcategory=None,
                    cuisine_group=cuisine_group,
                    source_order=order,
                )
            )
    return out


def extract_pork_beef(doc_path: Path) -> list[DishRow]:
    doc = Document(str(doc_path))
    category = None
    subcategory = None
    out: list[DishRow] = []
    order = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = getattr(p.style, "name", "") or ""
        if style == "Heading 1":
            if "pork" in text.lower():
                category = "pork"
            elif "beef" in text.lower():
                category = "beef"
            else:
                category = "pork/beef"
            subcategory = None
            continue
        if style == "Heading 2":
            subcategory = text
            continue
        if is_numbered_title_line(text):
            if category is None:
                category = "pork/beef"
            order += 1
            name = re.sub(r"^\s*\d+\.\s*", "", text).strip()
            out.append(
                DishRow(
                    display_name=name,
                    canonical_name=canonicalize_name(name),
                    category=category,
                    subcategory=subcategory,
                    cuisine_group=None,
                    source_order=order,
                )
            )
    return out


def extract_starters(doc_path: Path) -> list[DishRow]:
    doc = Document(str(doc_path))
    out: list[DishRow] = []
    order = 0
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = getattr(p.style, "name", "") or ""
        if style.startswith("Heading") and is_numbered_title_line(text):
            order += 1
            name = re.sub(r"^\s*\d+\.\s*", "", text).strip()
            out.append(
                DishRow(
                    display_name=name,
                    canonical_name=canonicalize_name(name),
                    category="starter",
                    subcategory=None,
                    cuisine_group=None,
                    source_order=order,
                )
            )
    return out


def extract_desserts(doc_path: Path) -> tuple[list[DishRow], dict[str, list[str]]]:
    doc = Document(str(doc_path))
    out: list[DishRow] = []
    ingredients_by_canonical: dict[str, list[str]] = {}
    order = 0
    for t in doc.tables:
        if not t.rows:
            continue
        header_cells = [c.text.strip() for c in t.rows[0].cells]
        if not header_cells:
            continue
        dish_cell = header_cells[0]
        if not is_numbered_title_line(dish_cell):
            continue
        order += 1
        name = re.sub(r"^\s*\d+\.\s*", "", dish_cell).strip()
        canonical = canonicalize_name(name)
        out.append(
            DishRow(
                display_name=name,
                canonical_name=canonical,
                category="dessert",
                subcategory=(header_cells[1] or None) if len(header_cells) > 1 else None,
                cuisine_group=None,
                source_order=order,
            )
        )
        items: list[str] = []
        for r in t.rows[1:]:
            for c in r.cells:
                txt = (c.text or "").strip()
                if not txt:
                    continue
                for line in txt.splitlines():
                    line = strip_bullet_prefix(line)
                    if line:
                        items.append(line)
        deduped = []
        seen = set()
        for it in items:
            key = it.lower()
            if key in seen:
                continue
            seen.add(key)
            deduped.append(it)
        ingredients_by_canonical[canonical] = deduped
    return out, ingredients_by_canonical


def init_db(conn: sqlite3.Connection, schema_path: Path) -> None:
    conn.executescript(schema_path.read_text(encoding="utf-8"))


def upsert_source_file(conn: sqlite3.Connection, path: str, category: str) -> int:
    ts = now_utc_iso()
    conn.execute(
        "INSERT INTO source_files(path, category, processed_at_utc) VALUES(?,?,?) "
        "ON CONFLICT(path) DO UPDATE SET category=excluded.category, processed_at_utc=excluded.processed_at_utc",
        (path, category, ts),
    )
    row = conn.execute("SELECT id FROM source_files WHERE path=?", (path,)).fetchone()
    return int(row[0])


def insert_dishes(
    conn: sqlite3.Connection,
    source_file_id: int,
    rows: list[DishRow],
    dessert_ingredients: dict[str, list[str]] | None = None,
) -> None:
    ts = now_utc_iso()
    for r in rows:
        conn.execute(
            "INSERT INTO dishes(canonical_name, display_name, category, subcategory, cuisine_group, source_file_id, source_order, created_at_utc, updated_at_utc) "
            "VALUES(?,?,?,?,?,?,?,?,?) "
            "ON CONFLICT(canonical_name, category) DO UPDATE SET "
            "display_name=excluded.display_name, subcategory=excluded.subcategory, cuisine_group=excluded.cuisine_group, "
            "source_file_id=excluded.source_file_id, source_order=excluded.source_order, updated_at_utc=excluded.updated_at_utc",
            (
                r.canonical_name,
                r.display_name,
                r.category,
                r.subcategory,
                r.cuisine_group,
                source_file_id,
                r.source_order,
                ts,
                ts,
            ),
        )
        dish_id = conn.execute(
            "SELECT id FROM dishes WHERE canonical_name=? AND category=?",
            (r.canonical_name, r.category),
        ).fetchone()[0]
        conn.execute(
            "INSERT INTO enrichment_queue(dish_id, status, updated_at_utc) VALUES(?, 'pending', ?) "
            "ON CONFLICT(dish_id) DO UPDATE SET status=excluded.status, updated_at_utc=excluded.updated_at_utc",
            (dish_id, ts),
        )

        if dessert_ingredients is not None and r.canonical_name in dessert_ingredients:
            conn.execute("DELETE FROM dish_ingredients WHERE dish_id=?", (dish_id,))
            for idx, ing in enumerate(dessert_ingredients[r.canonical_name], start=1):
                conn.execute(
                    "INSERT INTO dish_ingredients(dish_id, position, ingredient_text, created_at_utc) VALUES(?,?,?,?)",
                    (dish_id, idx, ing, ts),
                )


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--db", required=True, help="Path to sqlite db file")
    ap.add_argument("--schema", required=True, help="Path to schema.sql")
    ap.add_argument("--docx", required=True, help="Path to docx")
    ap.add_argument("--category", required=True, choices=["seafood", "dessert", "starter", "porkbeef"])
    args = ap.parse_args()

    db_path = Path(args.db)
    schema_path = Path(args.schema)
    doc_path = Path(args.docx)

    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    try:
        init_db(conn, schema_path)
        if args.category == "seafood":
            rows = extract_seafood(doc_path)
            sfid = upsert_source_file(conn, str(doc_path), "seafood")
            insert_dishes(conn, sfid, rows)
        elif args.category == "dessert":
            rows, ing = extract_desserts(doc_path)
            sfid = upsert_source_file(conn, str(doc_path), "dessert")
            insert_dishes(conn, sfid, rows, dessert_ingredients=ing)
        elif args.category == "starter":
            rows = extract_starters(doc_path)
            sfid = upsert_source_file(conn, str(doc_path), "starter")
            insert_dishes(conn, sfid, rows)
        elif args.category == "porkbeef":
            rows = extract_pork_beef(doc_path)
            sfid = upsert_source_file(conn, str(doc_path), "pork/beef")
            insert_dishes(conn, sfid, rows)
        conn.commit()
    finally:
        conn.close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

